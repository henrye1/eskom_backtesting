"""Convert the three audit documentation markdown files to Word (.docx) format.

Uses python-docx to produce professionally formatted documents with:
- Anchor Point Risk branding (navy headers)
- Proper heading hierarchy
- Formatted tables
- Bold/italic/code inline styles
- Page breaks between major sections
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── Branding colours ──────────────────────────────────────────────────
NAVY = RGBColor(0x2E, 0x75, 0xB6)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)

DOCS_DIR = Path(__file__).parent


def _set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def _add_formatted_run(paragraph, text: str):
    """Parse inline markdown (bold, italic, code, bold-italic) and add runs."""
    # Pattern handles: ***bold-italic***, **bold**, *italic*, `code`
    tokens = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('***') and token.endswith('***'):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            paragraph.add_run(token)


def _style_table(table):
    """Apply professional styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style header row
    if table.rows:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, '2E75B6')
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)
    # Style data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                _set_cell_shading(cell, 'F2F2F2')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert a markdown file to a formatted Word document."""
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK_GREY

    # ── Heading styles ──
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = NAVY
        hs.font.name = 'Calibri'

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # ── Code block toggle ──
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — write accumulated code
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table detection ──
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows (e.g. |---|---|)
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            _flush_table(doc, table_rows)
            table_rows = []
            in_table = False
            # Don't increment — reprocess this line

        # ── Horizontal rule ──
        if re.match(r'^-{3,}$', line.strip()):
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Headings ──
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            # Page break before H1 (except the very first)
            if level == 1 and i > 0:
                doc.add_page_break()
            p = doc.add_heading(level=level)
            _add_formatted_run(p, heading_text)
            i += 1
            continue

        # ── Bullet / numbered list ──
        bm = re.match(r'^(\s*)[-*]\s+(.*)', line)
        nm = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if bm:
            indent_level = len(bm.group(1)) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.27 + indent_level * 0.63)
            _add_formatted_run(p, bm.group(2))
            i += 1
            continue
        if nm:
            indent_level = len(nm.group(1)) // 2
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(1.27 + indent_level * 0.63)
            _add_formatted_run(p, nm.group(2))
            i += 1
            continue

        # ── Blank line ──
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ──
        p = doc.add_paragraph()
        _add_formatted_run(p, line)
        i += 1

    # Flush final table if file ends mid-table
    if in_table:
        _flush_table(doc, table_rows)

    doc.save(str(docx_path))
    print(f"  Saved: {docx_path.name}")


def _flush_table(doc, table_rows: list[list[str]]):
    """Write accumulated table rows to the document."""
    if not table_rows:
        return
    n_cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=n_cols)
    table.style = 'Table Grid'

    for ri, row_data in enumerate(table_rows):
        for ci, cell_text in enumerate(row_data):
            if ci < n_cols:
                cell = table.rows[ri].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_run(p, cell_text)

    _style_table(table)
    doc.add_paragraph()  # spacing after table


def main():
    files = [
        ('Dashboard_Metrics_Requirements.md', 'Dashboard_Metrics_Requirements.docx'),
        ('Dashboard_Metrics_RFD.md', 'Dashboard_Metrics_RFD.docx'),
        ('Dashboard_Metrics_Methodology.md', 'Dashboard_Metrics_Methodology.docx'),
    ]

    print("Converting markdown to Word...\n")
    for md_name, docx_name in files:
        md_path = DOCS_DIR / md_name
        docx_path = DOCS_DIR / docx_name
        if not md_path.exists():
            print(f"  SKIP: {md_name} not found")
            continue
        md_to_docx(md_path, docx_path)

    print("\nDone. All documents saved to:")
    print(f"  {DOCS_DIR}")


if __name__ == '__main__':
    main()
